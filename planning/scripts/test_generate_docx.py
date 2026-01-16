from docx import Document

def make_sample(path: str = "test_sample.docx"):
    doc = Document()
    doc.add_heading("Sample Document", level=1)
    p = doc.add_paragraph()
    p.add_run("This is a sample paragraph with ")
    r = p.add_run("bold text")
    r.bold = True
    p.add_run(" and ")
    r2 = p.add_run("italic text")
    r2.italic = True
    doc.add_paragraph("A second paragraph with a list:")
    doc.add_paragraph("First item", style="List Bullet")
    doc.add_paragraph("Second item", style="List Bullet")
    doc.save(path)

if __name__ == "__main__":
    make_sample()
